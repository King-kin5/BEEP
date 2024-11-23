from tkinter import filedialog, messagebox
from insert import *  # Import the insert_image function
from PIL import Image
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
import os
from tkinter import filedialog, messagebox
import fitz  # PyMuPDF
from fpdf import FPDF

            
def new_file(text_widget, root):
    if messagebox.askokcancel("New File", "Create new file? Unsaved changes will be lost."):
        text_widget.delete(1.0, "end")
        root.title("Beep (Text Editor) - New File")
        text_widget.image_paths = []  # Clear image paths

        # Destroy all existing tables
        if hasattr(text_widget, "tables"):
            for table in text_widget.tables:
                table.destroy()
            text_widget.tables = []


def save_file(text_widget, root):
    documents_path = os.path.expanduser("~/Documents")
    file_path = filedialog.asksaveasfilename(
        initialdir=documents_path,
        defaultextension=".docx",
        filetypes=[("Word documents", "*.docx"), ("PDF files", "*.pdf"), ("All files", "*.*")]
    )

    if file_path:
        try:
            if file_path.endswith(".docx"):
                doc = Document()
                
                # Save text and tables while retaining layout
                for i, widget_id in enumerate(text_widget.get("1.0", "end-1c").splitlines()):
                    print(f"Processing line {i}: {widget_id}")  # Debugging line
                    if widget_id.startswith("[TABLE]:"):  # Custom marker for tables
                        table_index = int(widget_id.split(":")[1])
                        table = text_widget.tables[table_index]
                        rows = len(table.cells)
                        cols = len(table.cells[0])
                        print(f"Saving table {table_index} with {rows} rows and {cols} columns")  # Debugging line
                        doc_table = doc.add_table(rows=rows, cols=cols)
                        for r in range(rows):
                            for c in range(cols):
                                cell_value = table.cells[r][c].get()
                                print(f"Saving cell [{r},{c}]: {cell_value}")  # Debugging line
                                doc_table.cell(r, c).text = cell_value
                    else:
                        # Add paragraphs with preserved spacing
                        print(f"Saving paragraph: {widget_id}")  # Debugging line
                        doc.add_paragraph(widget_id)

                doc.save(file_path)
                root.title(f"Beep (Text Editor) - {file_path}")
                print(f"File saved successfully as {file_path}")  # Debugging line

            messagebox.showinfo("Success", f"File saved successfully as {file_path}")

        except Exception as e:
            messagebox.showerror("Error", f"Could not save file: {str(e)}")
            print(f"Error saving file: {str(e)}")  # Debugging line


def open_file(text_widget, root):
    file_path = filedialog.askopenfilename(
        filetypes=[
            ("Word documents", "*.docx"),
            ("PDF files", "*.pdf"),
            ("All files", "*.*")
        ]
    )
    if file_path:
        try:
            text_widget.delete(1.0, "end")
            text_widget.tables = []  # Reset tables
            print("Cleared text widget and reset tables.")  # Debugging line

            if file_path.endswith(".docx"):
                doc = Document(file_path)

                for para in doc.paragraphs:
                    text_widget.insert("end", para.text + "\n")
                    print(f"Loaded paragraph: {para.text}")  # Debugging line

                for table_index, table in enumerate(doc.tables):
                    text_widget.insert("end", f"[TABLE]:{table_index}\n")  # Placeholder for table position
                    rows = len(table.rows)
                    cols = len(table.columns)
                    print(f"Loading table {table_index} with {rows} rows and {cols} columns")  # Debugging line
                    new_table = DraggableResizableTable(root, rows, cols, bg="white")
                    new_table.place(x=50, y=50)  # Adjust as needed

                    for r in range(rows):
                        for c in range(cols):
                            cell_value = table.cell(r, c).text
                            print(f"Loading cell [{r},{c}]: {cell_value}")  # Debugging line
                            new_table.cells[r][c].insert(0, cell_value)

                    text_widget.tables.append(new_table)

                root.title(f"Beep (Text Editor) - {file_path}")
                messagebox.showinfo("Success", "Word document opened successfully!")
                print("File opened successfully.")  # Debugging line

        except Exception as e:
            messagebox.showerror("Error", f"Could not open file: {str(e)}")
            print(f"Error opening file: {str(e)}")  # Debugging line